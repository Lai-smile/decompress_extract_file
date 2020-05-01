# Created by LYN at 2019/2/19

"""
Feature: #Enter feature name here
# Enter feature description here

Scenario: #Enter scenario name here
# Enter steps here

Test File Location: # Enter

"""
import os

from bs4 import BeautifulSoup
import re
import docx
from docx.shared import RGBColor
import pickle
from docx.document import Document
from zipfile import ZipFile
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml.shape import CT_Picture
from docx.shape import InlineShape
from get_content_from_word import doc2docx
import pythoncom
from selected_value_extract import full_check_flag_list, extract_selected_value
from txt_utilities import any_item_be_contain
from log.logger import logger

# def get_hyperlink(xml_path):
#     bs = BeautifulSoup(open(xml_path, encoding="utf-8"), "html.parser")
#     for node in bs.find_all("w:t"):
#

def get_highlighted_checkbox(xml_path):
    bs = BeautifulSoup(open(xml_path, encoding="utf-8"), "html.parser")
    # print(bs.prettify())
    selected = []
    for node in bs.find_all("w:highlight"):
        np = node.parent
        ns = np.next_sibling
        if ns.text == "□":
            text_node = ns.parent.next_sibling
            # print(text_node)
            text = text_node.find("w:t").text
            if text.strip():
                # if "□" in text:
                #     text_list=re.findall("^(.*?)□",text)
                #     text=text_list[0].strip()

                selected.append(text)
            else:
                text_node_next = text_node.next_sibling
                text_next = text_node_next.find("w:t").text
                selected.append(text_next)
    return selected


def get_form_checkbox(xml_path):
    bs = BeautifulSoup(open(xml_path, encoding="utf-8"), "html.parser")
    for node in bs.find_all("w:checkbox"):
        #         <w:default w:val="1"/>
        if node.find(attrs={"w:val": "1"}):
            print()


def add_mark2txt(txtfile, selected_list):
    with open(txtfile, "r", encoding="utf-8") as f:
        f.seek(0)
        contents = f.read()
        print("*****************************************************8", contents)

    # print(contents)
    for t in selected_list:
        # pattern_parentheses = re.compile(u'\\(|\\)')
        # t = pattern_parentheses.sub(' ', t)
        t = t.split("(")[0]
        pattern = r"□{1}\s*" + t

        selected_text = re.findall(pattern, contents)
        print(t, "******is found here", selected_text)
        if len(selected_text) == 1:
            contents = re.sub(
                selected_text[0],
                "woshiyigebeixuanzhongdekuang" + selected_text[0],
                contents)
            print(contents)
    marked_txt = txtfile + "-marked.txt"
    with open(marked_txt, 'w', encoding="utf-8") as f:
        f.write(contents)

    return marked_txt


def iter_unique_cells(table):
    prior_tc = None
    tr = 1
    for row in table.rows:

        tc = 0
        for cell in row.cells:
            this_tc = cell._tc
            if this_tc is prior_tc:
                continue
            tc += 1
            prior_tc = this_tc
            yield tr, tc, cell
        tr += 1


def clean_pos(data):
    cleaned = data[:]
    for i in range(len(data)):
        print(data[i])
        if i > 0 and (data[i][1] - data[i - 1][1]) > 0:
            cleaned[i - 1][1] += 1
            cleaned[i - 1][2] = 1
            cleaned[i][2] += 1

    return cleaned


def get_tables(docxfile):
    doc = docx.Document(docxfile)
    t = 0
    res = []
    for tbl in doc.tables:
        t += 1
        for r, c, cell in iter_unique_cells(tbl):
            cell_text = ""
            for p in cell.paragraphs:
                for run in p.runs:
                    cell_text += run.text
                    if run.font.highlight_color:
                        cell_text += str(run.font.highlight_color)

            res.append([t, r, c, cell_text])
    cleaned_data = clean_pos(res)

    return cleaned_data


def get_hyperlink(filename):
    docxf = doc2docx(filename)
    document = ZipFile(docxf)
    xml = document.read("word/document.xml")
    wordObj = BeautifulSoup(xml.decode("utf-8"))
    hyperlink = []
    for links in wordObj.find_all("w:hyperlink"):
        for child in links.find_all("w:t"):
            child = child.string
            hyperlink.append(child)
    hyperlink = ''.join(str(t) for t in hyperlink)
    return hyperlink


def get_table_text(tbl, num):
    res = []
    for r, c, cell in iter_unique_cells(tbl):
        cell_text = ""
        for p in cell.paragraphs:
            for run in p.runs:
                cell_text += run.text
                # if run.font.highlight_color:
                #     cell_text += str(run.font.highlight_color)
                # if cell_text.endswith(' '):
                # print('True')
            # cell_text += get_hyperlink(filename)

        res.append(('t{}'.format(num), r, c, r, c, cell_text))
    return res


def get_table_text_new(tbl, num):
    res = []
    tr = 1
    num_flag = 1
    for row in tbl.rows:
        tc = 1
        tc_list = row._tr.tc_lst
        for tc_item in tc_list:
            #     tc_item.
            # for cell in row.cells:
            #     this_tc = cell._tc
            this_tc = tc_item
            # cell_text = cell.text
            # this_plist = this_tc.p_lst
            this_plist = list(this_tc.iter_block_items())
            cell_text = ''
            for p_index, p in enumerate(this_plist):
                # added by zhoulong for embedding table
                if isinstance(p, CT_Tbl):
                    for child_table_row in p.tr_lst:
                        tc = 1
                        for child_table_cell in child_table_row.tc_lst:
                            # print(child_table_cell.p_lst)

                            for cur_row_str in child_table_cell.p_lst:
                                cur_cell_str = ''
                                if cur_row_str.r_lst:
                                    for cur_row_str_cur_r in cur_row_str.r_lst:
                                        cur_cell_str += cur_row_str_cur_r.text
                                    cur_cell_str += ' '
                                # print(('t{}'.format(num), tr, tc, tr, tc, cur_cell_str))
                                res.append(('t{}'.format(num), tr, tc, tr, tc, cur_cell_str.strip()))
                            tc += 1
                        tr += 1
                    continue
                # deal with embedding table
                this_pPr = p.pPr
                if this_pPr and this_pPr.numPr:
                    if p.r_lst:
                        for index, r in enumerate(p.r_lst):
                            if index == 0:
                                cell_text += str(p_index + 1) + ') ' + r.text
                            else:
                                cell_text += r.text
                    else:
                        cell_text = str(num_flag)
                        num_flag += 1
                    continue
                if 'checkBox' in xml_to_dict(p.xml):
                    checkBox_list = get_xml_object(p.xml, 'default')
                    checkBox_t_list = get_xml_object(p.xml, 't')
                    for checkbox_index, checkbox in enumerate(checkBox_list):
                        if checkbox.attrs['w:val'] == '1':
                            cell_text = checkBox_t_list[checkbox_index].text
                            # res.append(('t{}'.format(num), tr, tc, tr, tc, cell_text))
                    continue
                elif 'hyperlink' in xml_to_dict(p.xml):
                    hyperlink_t_list = get_xml_object(p.xml, 't')
                    for hyperlink_t_list_item in hyperlink_t_list:
                        cell_text += hyperlink_t_list_item.text
                    # res.append(('t{}'.format(num), tr, tc, tr, tc, cell_text))
                    continue
                #elif 'smartTag' in xml_to_dict(p.xml):
                #    smartTagXML = get_xml_object(p.xml, 'smartTag')
                #    for smartTag in smartTagXML:
                #        cell_text += smartTag.text.strip("\n")
                #    cell_text = cell_text.replace(' ', '').replace('\n', '').replace('\t', '')
                #    continue
                else:
                    this_rlist = p.r_lst
                    if this_rlist:
                        for r in this_rlist:
                            txt = r.text
                            cell_text += txt
                        if cell_text and any_item_be_contain(full_check_flag_list, cell_text):
                            cell_text = extract_selected_value(cell_text)
                        cell_text += ' '
            print(('t{}'.format(num), tr, tc, tr, tc, cell_text))
            res.append(('t{}'.format(num), tr, tc, tr, tc, cell_text.strip()))
            tc += 1
        tr += 1
    # print(res)
    return res


def xml_to_dict(xml_data):
    """
    xml转换为字典
    :param xml_data:
    :return:
    """
    soup = BeautifulSoup(xml_data, features='xml')
    xml = soup.is_xml
    if not xml:
        return {}
    # 将 XML 数据转化为 Dict
    data = dict([(item.name, item.text) for item in soup.find_all()])
    return data


def get_xml_object(xml_data, key):
    """
        xml object
        :param xml_data:
        :return:
    """
    soup = BeautifulSoup(xml_data, features='xml')
    xml = soup.is_xml
    if not xml:
        return None
    checkbox_list = []
    for item in soup.find_all():
        if item.name == key:
            checkbox_list.append(item)
    return checkbox_list


def get_text(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        for run in para.runs:
            if run.font.highlight_color:
                print(run.text, run.font.highlight_color)
    return fullText


def iter_block_items(parent):
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("sth's wrong")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Picture):
            yield InlineShape(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

            # table = Table(child, parent)
            # for row in table.rows:
            #     for cell in row.cells:
            #         yield from iter_block_items(cell)


def pre_story(input_file, un_match_tokens=None):
    # print(un_match_tokens)
    un_match_tokens = un_match_tokens.split('$zl$')
    try:
        import os
        result_file = input_file[:input_file.rfind('.')] + '_new' + input_file[input_file.rfind('.'):]
        word_obj = docx.Document(input_file)

        for cur_table_index in range(len(word_obj.tables)):
            cur_table = word_obj.tables[cur_table_index]
            for cur_row_index_cur_table in range(len(cur_table.rows)):
                for cur_col_index_cur_row in range(len(cur_table.columns)):
                    # # 得到当前单元格
                    cur_cell = cur_table.cell(cur_row_index_cur_table, cur_col_index_cur_row)
                    cur_tc_list = cur_cell._tc
                    cur_cell_text_list = []

                    for cur_tc_index in range(len(cur_tc_list) - 1, -1, -1):
                        # 初始化变量
                        cur_tc = cur_tc_list[cur_tc_index]
                        delete_and_add_element = -1
                        delete_and_add_element += 1
                        cur_tc_text = ''

                        if isinstance(cur_tc, CT_P):
                            # # 如果是文本段落的话, 判断段落是否空，空的话跳过，不空的话，删掉重新加
                            # print(cur_tc, cur_tc.r_lst)

                            for cur_r_in_cur_tc_r_list in cur_tc.r_lst:
                                cur_tc_text += cur_r_in_cur_tc_r_list.text

                            # print(cur_tc, cur_tc_text, bool(cur_tc_text and not cur_tc_text.isspace()))

                            if bool(cur_tc_text and not cur_tc_text.isspace()):
                                # 删掉这个，并增加一个段落。
                                delete_and_add_element = 1
                                cur_cell_text_list.insert(0, cur_tc_text)
                            else:
                                delete_and_add_element = 0

                        elif isinstance(cur_tc, CT_Tbl):
                            delete_and_add_element = 0
                            child_table_cur_tc = cur_tc
                            for child_table_row in child_table_cur_tc.tr_lst:
                                for child_table_cell in child_table_row.tc_lst:
                                    # 嵌套表的当前单元格 child_table_cell
                                    block_list_child_table_cell = list(child_table_cell.iter_block_items())
                                    cur_cell_child_table_text_list = []
                                    for cur_block_child_table_cell_index in range(len(block_list_child_table_cell) - 1,
                                                                                  -1, -1):
                                        # 初始化嵌套表里的变量。
                                        cur_block_child_table_cell = \
                                            block_list_child_table_cell[cur_block_child_table_cell_index]
                                        delete_and_add_element_child_table = -1
                                        delete_and_add_element_child_table += 1
                                        cur_tc_text_child_table = ''

                                        if isinstance(cur_block_child_table_cell, CT_P):
                                            # # 如果是文本段落的话, 判断段落是否空，空的话跳过，不空的话，删掉重新加

                                            for cur_r_in_cur_tc_r_list in cur_block_child_table_cell.r_lst:
                                                cur_tc_text_child_table += cur_r_in_cur_tc_r_list.text
                                            # import pdb
                                            # pdb.set_trace()
                                            # print(cur_tc_text_child_table, bool(cur_tc_text_child_table and not cur_tc_text_child_table.isspace()))
                                            if bool(cur_tc_text_child_table and not cur_tc_text_child_table.isspace()):
                                                # 删掉这个，并增加一个段落。
                                                delete_and_add_element_child_table = 1
                                                cur_cell_child_table_text_list.insert(0, cur_tc_text_child_table)
                                            else:
                                                delete_and_add_element_child_table = 0
                                        else:
                                            # # 如果不是CT_P，就跳过
                                            delete_and_add_element_child_table = 0

                                        if delete_and_add_element_child_table == 1:
                                            for p_zl in block_list_child_table_cell[
                                                cur_block_child_table_cell_index].r_lst:
                                                p_zl.text = ''
                                            del block_list_child_table_cell[cur_block_child_table_cell_index]

                                    # # 在嵌套表中增加段落，改颜色
                                    for cur_add_text_index, cur_add_text in enumerate(cur_cell_child_table_text_list):
                                        cell_buffer = _Cell(child_table_cell, child_table_cur_tc)
                                        run = cell_buffer.paragraphs[-1].add_run(cur_add_text)
                                        # if cur_add_text_index == 0:
                                        #     run = cell_buffer.paragraphs[-1].add_run(cur_add_text)
                                        # else:
                                        #     run = cell_buffer.paragraphs[-1].add_run(cur_add_text)
                                        run.font.name = '宋体'
                                        run.font.size = 140000
                                        # run.font.color.rgb = RGBColor(255, 0, 0)
                                        if not cur_add_text.isspace():
                                            run.font.highlight_color = 4
                                        if un_match_tokens is not None and cur_add_text in un_match_tokens:
                                            run.font.highlight_color = 7
                                            un_match_tokens.remove(cur_add_text)
                        else:
                            # # 其他情况
                            delete_and_add_element = 0

                        if delete_and_add_element == 1:
                            del cur_tc_list[cur_tc_index]

                    for cur_add_text_index, cur_add_text in enumerate(cur_cell_text_list):
                        if cur_add_text_index == 0:
                            run = cur_cell.add_paragraph().add_run(cur_add_text)
                        else:
                            run = cur_cell.paragraphs[-1].add_run(cur_add_text)
                        run.font.name = '宋体'
                        run.font.size = 140001
                        # run.font.color.rgb = RGBColor(255, 0, 0)
                        if not cur_add_text.isspace():
                            run.font.highlight_color = 4
                        if un_match_tokens is not None and cur_add_text in un_match_tokens:
                            run.font.highlight_color = 7
                            un_match_tokens.remove(cur_add_text)

        for i in range(len(word_obj.paragraphs)):
            cur_p = word_obj.paragraphs[i]
            cur_p_text = cur_p.text
            cur_p.text = ''
            run = cur_p.add_run(cur_p_text)
            if not cur_p_text.isspace():
                run.font.highlight_color = 4
            #print('zlzlzl', cur_p_text, un_match_tokens)
            if un_match_tokens is not None and cur_p_text in un_match_tokens:
                run.font.highlight_color = 7
                un_match_tokens.remove(cur_p_text)

        word_obj.save(result_file)
    except Exception as e:
        result_file = input_file
    return result_file


def story(docxfile):
    # Added by zhoulong
    # To solve some problem
    # docxfile = pre_story(docxfile, un_match_tokens=None)
    # print('zl', docxfile)
    doc = docx.Document(docxfile)
    res = []
    i, num = 0, 0
    for block in iter_block_items(doc):
        i += 1
        if isinstance(block, Paragraph):
            if block.text:
                res.append((0, i, 0, 0, 0, ["Text", block.text]))
            # print("para*********", block.text)
        elif isinstance(block, Table):
            num += 1
            tbl_text = get_table_text_new(block, num)
            res.append((0, i, 0, 0, 0, ["Table", tbl_text]))
            # print("tbl**********", tbl_text)
    print(res)
    return res


def get_content_from_all_version(filename):
    try:
        # os.system("taskkill /f /im word.exe /t")
        pythoncom.CoInitialize()
        if filename.endswith('doc'):
            docxfile = doc2docx(filename)
        elif filename.endswith("docx"):
            docxfile = filename
        else:
            return "Not Word Document"
        return story(docxfile)
    except Exception as e:
        logger.error("Failed to get content.", exc_info=True)
        return 'The file may be damaged. Please save it as .doc/.docx file and upload it again.文件可能已经损坏，请另存为doc或docx文件并重新上传。'


def get_checkBox(filename):
    if filename.endswith('doc'):
        docxfile = doc2docx(filename)
    elif filename.endswith("docx"):
        docxfile = filename
    doc = docx.Document(docxfile)
    doc_elm = doc._element
    checkBoxes = doc_elm.xpath('w:checkBox')

    for checkBox in checkBoxes:
        print('checkBox value is %s' % checkBox)


if __name__ == '__main__':
    # filename = r"C:\Users\Administrator\Desktop\ghv.doc"
    filename = r"C:\Users\Administrator\Desktop\dsg.doc"
    txtfile = r"D:/IBM项目/doc/─┌╡Ñ─ú░σ--S7BL.doc"
    # get_checkBox(filename)
    # docx= r"D:/IBM项目/doc/─┌╡Ñ─ú░σ--S7BL.doc"
    # picklefile = r"C:\Users\LYN\Desktop\word\process.pickle"
    # st = get_highlighted_checkbox(xmlfile)
    # add_mark2txt(txtfile, st)
    # get_form_checkbox(xmlfile)
    # print(res)
    docxfile = doc2docx(filename)
    cleaned_data = get_tables(docxfile)
    print(cleaned_data)
    # with open(picklefile, "wb") as f:
    #     pickle.dump(cleaned_data, f)
    # text=get_text(docxfile)
    # print(text)

    story = get_content_from_all_version(filename)
    print(story)
